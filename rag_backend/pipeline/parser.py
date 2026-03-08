"""
pipeline/parser.py
PDF / DOCX / Markdown -> 结构化 ParsedDocument

职责划分：
1. 统一 dispatcher：根据后缀选择解析器
2. PDF：文本 + 表格 + 图片占位，失败时 fallback 到 PyMuPDF
3. DOCX：按正文块顺序提取段落和表格，并保留标题层级
4. Markdown：按标题切分为虚拟页
5. 后处理：删除模板噪声、清理控制字符
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable, Optional

import fitz
import pdfplumber
from loguru import logger


@dataclass
class ParsedPage:
    page_num: int
    text: str


@dataclass
class ParsedDocument:
    file_path: str
    file_name: str
    pages: list[ParsedPage] = field(default_factory=list)
    markdown_content: Optional[str] = None

    @property
    def full_text(self) -> str:
        if self.markdown_content:
            return self.markdown_content
        return "\n\n".join(page.text for page in self.pages if page.text.strip())


_HEADING_MIN_SIZE = 11.0
_DOCX_VIRTUAL_PAGE_LINE_LIMIT = 60
_NOISE_PATTERNS = [
    re.compile(r"^[\*_~>]{0,3}\s*本条应.{0,60}$"),
    re.compile(r"^[\*_~>]{0,3}\s*说明：.*$"),
    re.compile(r"^[\*_~>]{0,3}\s*最终文档请删除.*$"),
    re.compile(r"^[\*_~>]{0,3}\s*\[?TOC\]?$", re.IGNORECASE),
]


def parse_document(path: str | Path) -> ParsedDocument:
    """根据文件后缀自动选择解析器。"""
    file_path = Path(path)
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return parse_pdf(file_path)
    if suffix == ".docx":
        return parse_docx(file_path)
    if suffix in {".md", ".markdown"}:
        return parse_md(file_path)

    raise ValueError(f"不支持的文件类型：{suffix}，支持 .pdf / .docx / .md")


def parse_pdf(path: str | Path) -> ParsedDocument:
    """解析 PDF，尽可能保留标题、表格和图片提示信息。"""
    file_path = Path(path)
    document = ParsedDocument(file_path=str(file_path), file_name=file_path.name)
    logger.info(f"开始解析 PDF：{file_path.name}")

    try:
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                raw_text = _extract_pdf_page(page, file_path.name, page_num)
                cleaned_text = _remove_template_noise(raw_text)
                if cleaned_text.strip():
                    document.pages.append(ParsedPage(page_num=page_num, text=cleaned_text))
    except Exception as exc:
        logger.warning(
            f"pdfplumber 整体解析失败，切换到 PyMuPDF（file={file_path.name}）：{type(exc).__name__}: {exc}"
        )
        _parse_pdf_with_pymupdf(file_path, document, exc)

    logger.info(f"PDF 解析完成：{file_path.name}，共 {len(document.pages)} 页有效内容")
    return document


def parse_docx(path: str | Path) -> ParsedDocument:
    """解析 DOCX，按文档块顺序保留标题和表格。"""
    from docx import Document as DocxDocument
    from docx.document import Document as DocxDocumentType
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    file_path = Path(path)
    document = ParsedDocument(file_path=str(file_path), file_name=file_path.name)
    logger.info(f"开始解析 DOCX：{file_path.name}")

    try:
        docx_document = DocxDocument(str(file_path))
    except Exception as exc:
        logger.exception(f"DOCX 解析失败（file={file_path.name}）")
        raise RuntimeError(f"DOCX 解析失败：{type(exc).__name__}: {exc}") from exc

    current_lines: list[str] = []
    page_num = 1

    def flush_page() -> None:
        nonlocal page_num
        content = _remove_template_noise("\n".join(current_lines).strip())
        if content:
            document.pages.append(ParsedPage(page_num=page_num, text=content))
            page_num += 1
        current_lines.clear()

    def iter_blocks(doc: DocxDocumentType) -> Iterable[Paragraph | Table]:
        body = doc.element.body
        for child in body.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, doc)
            elif isinstance(child, CT_Tbl):
                yield Table(child, doc)

    for block in iter_blocks(docx_document):
        if isinstance(block, Paragraph):
            paragraph_text = _clean_line(block.text)
            if not paragraph_text:
                continue

            heading_level = _heading_level_from_style(block.style.name if block.style else "")
            if heading_level is not None:
                flush_page()
                current_lines.append(f"{'#' * heading_level} {paragraph_text}")
            else:
                current_lines.append(paragraph_text)
        else:
            table_markdown = _docx_table_to_markdown(block)
            if table_markdown:
                current_lines.append(table_markdown)

        if len(current_lines) >= _DOCX_VIRTUAL_PAGE_LINE_LIMIT:
            flush_page()

    flush_page()
    logger.info(f"DOCX 解析完成：{file_path.name}，共 {len(document.pages)} 页有效内容")
    return document


def parse_md(path: str | Path) -> ParsedDocument:
    """解析 Markdown，按一级/二级标题切分为虚拟页。"""
    file_path = Path(path)
    document = ParsedDocument(file_path=str(file_path), file_name=file_path.name)
    logger.info(f"开始解析 Markdown：{file_path.name}")

    try:
        text = file_path.read_text(encoding="utf-8", errors="replace")
    except Exception as exc:
        logger.exception(f"Markdown 读取失败（file={file_path.name}）")
        raise RuntimeError(f"Markdown 读取失败：{type(exc).__name__}: {exc}") from exc

    cleaned_text = _remove_template_noise(text)
    sections = _split_markdown_sections(cleaned_text)
    for page_num, section in enumerate(sections, 1):
        if section:
            document.pages.append(ParsedPage(page_num=page_num, text=section))

    logger.info(f"Markdown 解析完成：{file_path.name}，共 {len(document.pages)} 页有效内容")
    return document


def _extract_pdf_page(page: pdfplumber.page.Page, file_name: str, page_num: int) -> str:
    parts: list[str] = []

    try:
        text = _parse_page_pdfplumber(page)
    except Exception as exc:
        logger.warning(
            f"pdfplumber 精细解析失败，回退基础提取（file={file_name}, page={page_num}）："
            f"{type(exc).__name__}: {exc}"
        )
        text = page.extract_text() or ""

    text = text.strip()
    if text:
        parts.append(text)

    table_markdown = _pdf_tables_to_markdown(page.extract_tables())
    if table_markdown:
        parts.append(table_markdown)

    image_note = _pdf_image_note(page)
    if image_note:
        parts.append(image_note)

    return "\n\n".join(parts)


def _parse_pdf_with_pymupdf(file_path: Path, document: ParsedDocument, original_exc: Exception) -> None:
    try:
        fitz_doc = fitz.open(str(file_path))
        try:
            for page_num, page in enumerate(fitz_doc, 1):
                raw_text = _parse_page_pymupdf(page)
                cleaned_text = _remove_template_noise(raw_text)
                if cleaned_text.strip():
                    document.pages.append(ParsedPage(page_num=page_num, text=cleaned_text))
        finally:
            fitz_doc.close()
    except Exception as fallback_exc:
        logger.exception(f"PyMuPDF fallback 解析失败（file={file_path.name}）")
        raise RuntimeError(
            f"PDF 解析失败（pdfplumber={type(original_exc).__name__}: {original_exc}; "
            f"pymupdf={type(fallback_exc).__name__}: {fallback_exc}）"
        ) from fallback_exc


def _size_to_heading(size: float, is_bold: bool) -> Optional[str]:
    if size >= 18:
        return "# "
    if size >= 15:
        return "## "
    if size >= 13:
        return "### "
    if size >= _HEADING_MIN_SIZE and is_bold:
        return "#### "
    return None


def _heading_level_from_style(style_name: str) -> int | None:
    lowered = style_name.lower().strip()
    if "heading 1" in lowered or "标题 1" in style_name:
        return 1
    if "heading 2" in lowered or "标题 2" in style_name:
        return 2
    if "heading 3" in lowered or "标题 3" in style_name:
        return 3
    if "heading" in lowered or "标题" in style_name:
        return 4
    return None


def _clean_line(line: str) -> str:
    line = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", line)
    line = re.sub(r" {3,}", "  ", line)
    return line.strip()


def _parse_page_pdfplumber(page: pdfplumber.page.Page) -> str:
    words = page.extract_words(extra_attrs=["fontname", "size"], keep_blank_chars=False)
    if not words:
        return page.extract_text() or ""

    lines: list[str] = []
    current_line_words: list[dict] = []
    current_y: float | None = None

    for word in words:
        y = round(float(word.get("top", 0)), 1)
        if current_y is None or abs(y - current_y) < 3:
            current_y = y
            current_line_words.append(word)
            continue

        lines.append(_flush_pdf_line(current_line_words))
        current_line_words = [word]
        current_y = y

    if current_line_words:
        lines.append(_flush_pdf_line(current_line_words))

    return "\n".join(lines)


def _flush_pdf_line(words: list[dict]) -> str:
    sizes = [float(word.get("size", 10)) for word in words]
    fonts = [word.get("fontname", "") for word in words]
    text = _clean_line(" ".join(word["text"] for word in words))

    average_size = sum(sizes) / len(sizes) if sizes else 10
    is_bold = any("Bold" in font or "bold" in font for font in fonts)
    prefix = _size_to_heading(average_size, is_bold)
    return f"{prefix}{text}" if prefix else text


def _parse_page_pymupdf(page: fitz.Page) -> str:
    lines: list[str] = []
    for block in page.get_text("dict").get("blocks", []):
        for line in block.get("lines", []):
            spans = line.get("spans", [])
            texts = [span["text"] for span in spans if span.get("text", "").strip()]
            if not texts:
                continue

            full_line = _clean_line(" ".join(texts))
            max_size = max((span.get("size", 10) for span in spans), default=10)
            is_bold = any("bold" in span.get("font", "").lower() for span in spans)
            prefix = _size_to_heading(max_size, is_bold)
            lines.append(f"{prefix}{full_line}" if prefix else full_line)
    return "\n".join(lines)


def _pdf_tables_to_markdown(tables: list[list[list[str | None]]] | None) -> str:
    if not tables:
        return ""

    rendered_tables: list[str] = []
    for table in tables:
        normalized_rows = [_normalize_table_row(row) for row in table if row and any(cell for cell in row)]
        if not normalized_rows:
            continue
        rendered_tables.append(_rows_to_markdown_table(normalized_rows))

    return "\n\n".join(rendered_tables)


def _pdf_image_note(page: pdfplumber.page.Page) -> str:
    image_count = len(page.images or [])
    if image_count <= 0:
        return ""
    return f"> [图例：本页包含 {image_count} 张图片/架构图，详情请参考原始文档。]"


def _docx_table_to_markdown(table) -> str:
    rows = []
    for row in table.rows:
        normalized = [_clean_line(cell.text) for cell in row.cells]
        if any(normalized):
            rows.append(normalized)

    if not rows:
        return ""
    return _rows_to_markdown_table(rows)


def _rows_to_markdown_table(rows: list[list[str]]) -> str:
    if not rows:
        return ""

    width = max(len(row) for row in rows)
    padded_rows = [row + [""] * (width - len(row)) for row in rows]
    header = padded_rows[0]
    body = padded_rows[1:] or [[""] * width]

    markdown_lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join("---" for _ in range(width)) + " |",
    ]
    markdown_lines.extend("| " + " | ".join(row) + " |" for row in body)
    return "\n".join(markdown_lines)


def _normalize_table_row(row: list[str | None]) -> list[str]:
    return [str(cell).replace("\n", " ").strip() if cell else "" for cell in row]


def _remove_template_noise(text: str) -> str:
    cleaned_lines = []
    for line in text.splitlines():
        stripped = line.strip()
        if any(pattern.match(stripped) for pattern in _NOISE_PATTERNS):
            continue
        cleaned_lines.append(line)
    return "\n".join(cleaned_lines)


def _split_markdown_sections(text: str) -> list[str]:
    sections: list[str] = []
    current: list[str] = []

    for line in text.splitlines():
        if re.match(r"^#{1,2} ", line) and current:
            sections.append("\n".join(current).strip())
            current = [line]
        else:
            current.append(line)

    if current:
        sections.append("\n".join(current).strip())
    return [section for section in sections if section]